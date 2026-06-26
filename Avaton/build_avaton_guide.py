from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK

OUT = r"D:\Softwheres\Avaton\Avaton_Complete_Guide.docx"

NAVY = "20243A"
PURPLE = "7559E8"
PURPLE_DARK = "5941BD"
INK = "242735"
MUTED = "6F7482"
LIGHT = "F3F1FC"
BLUE_LIGHT = "EAF2FA"
GOLD_LIGHT = "FFF5D9"
RED_LIGHT = "FDECEF"
GREEN_LIGHT = "E9F7EF"
WHITE = "FFFFFF"
BORDER = "D9DCE5"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.82)
sec.bottom_margin = Inches(0.78)
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)
sec.header_distance = Inches(0.38)
sec.footer_distance = Inches(0.38)

def set_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18

for style_name, size, color, before, after in [
    ("Heading 1", 16, PURPLE_DARK, 16, 8),
    ("Heading 2", 13, PURPLE_DARK, 12, 6),
    ("Heading 3", 11.5, NAVY, 9, 4),
]:
    st = styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for name in ["Guide Kicker", "Guide Callout Title"]:
    if name not in styles:
        styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
styles["Guide Kicker"].font.name = "Calibri"
styles["Guide Kicker"].font.size = Pt(9)
styles["Guide Kicker"].font.bold = True
styles["Guide Kicker"].font.color.rgb = RGBColor.from_string(PURPLE)
styles["Guide Kicker"].paragraph_format.space_after = Pt(8)
styles["Guide Callout Title"].font.name = "Calibri"
styles["Guide Callout Title"].font.size = Pt(10.5)
styles["Guide Callout Title"].font.bold = True
styles["Guide Callout Title"].font.color.rgb = RGBColor.from_string(NAVY)
styles["Guide Callout Title"].paragraph_format.space_after = Pt(2)

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc.get_or_add_tcPr()
    tc_mar = tc.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc.append(tc_mar)
    for tag, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def border_bottom(paragraph, color=PURPLE, size=12):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)

header = sec.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = hp.add_run("AVATON")
set_font(r, size=9, color=PURPLE, bold=True)
r = hp.add_run("  |  Complete Functionality, Storage & Backup Guide")
set_font(r, size=9, color=MUTED)

footer = sec.footer
fp = footer.paragraphs[0]
add_page_field(fp)

def para(text="", bold_start=None, color=INK, size=10.5, align=None, after=6, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if bold_start and text.startswith(bold_start):
        a = p.add_run(bold_start)
        set_font(a, size=size, color=color, bold=True)
        b = p.add_run(text[len(bold_start):])
        set_font(b, size=size, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_font(run, size=size, color=color, italic=italic)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.38 + level * 0.24)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(text), size=10.3)
    return p

def number(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.42)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run(text), size=10.3)
    return p

def heading(text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")

def callout(title, body, fill=LIGHT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.style = styles["Guide Callout Title"]
    p.add_run(title)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_font(p2.add_run(body), size=10.1, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def key_value_table(rows, widths=(2400, 6960), header=None):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    if header:
        row = table.add_row()
        row.cells[0].text = header[0]
        row.cells[1].text = header[1]
        for cell in row.cells:
            shade(cell, PURPLE_DARK)
            for run in cell.paragraphs[0].runs:
                set_font(run, size=10, color=WHITE, bold=True)
    for label, detail in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        shade(cells[0], LIGHT)
        for run in cells[0].paragraphs[0].runs:
            set_font(run, size=9.8, color=NAVY, bold=True)
        for run in cells[1].paragraphs[0].runs:
            set_font(run, size=9.8, color=INK)
    set_table_geometry(table, list(widths))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(55)
p = doc.add_paragraph(style="Guide Kicker")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("PERSONAL VAULT REFERENCE GUIDE")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run("Avaton")
set_font(r, name="Calibri", size=34, color=NAVY, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(22)
r = p.add_run("Complete Functionality, Storage, Backup & Limitations Guide")
set_font(r, size=16, color=PURPLE_DARK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(46)
r = p.add_run("A practical manual for using Avaton safely, understanding where your data lives, and recovering everything after a browser or computer reset.")
set_font(r, size=11.5, color=MUTED, italic=True)
callout(
    "The most important rule",
    "Avaton's working copy lives inside your browser profile. A backup is a separate ZIP file. Keep at least one tested backup on a USB drive or another computer before clearing browser data, reinstalling Windows, changing browsers, or deleting the laptop.",
    GOLD_LIGHT,
)
para("Document version: June 24, 2026", color=MUTED, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("Applies to the current HTML/CSS/JavaScript build in this workspace.", color=MUTED, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# Contents
heading("Quick navigation", 1)
key_value_table([
    ("1", "What Avaton is and what it can do"),
    ("2", "Password, lock screen, hint, and theme"),
    ("3", "Scripts, nested scripts, folders, and attachments"),
    ("4", "File upload, download, and preview behavior"),
    ("5", "Cells: the Excel-like workbook"),
    ("6", "Calendar and countdown behavior"),
    ("7", "Exactly where data is stored"),
    ("8", "How full and single-script backups work"),
    ("9", "What happens to current data during restore"),
    ("10", "Limitations and practical safety rules"),
    ("11", "Recommended recovery workflow"),
], header=("Section", "Topic"))
callout("At a glance", "Avaton is local-first. It does not send scripts or attachments to a server. Privacy and portability therefore depend on the browser profile and the backups you create.", BLUE_LIGHT)

heading("1. What Avaton is", 1)
para("Avaton is a private, browser-based personal vault built with HTML, CSS, and JavaScript. It combines notes, nested scripts, folders, file storage, a countdown calendar, a spreadsheet-style Cells workbook, file previews, password protection, and portable backup/restore.")
heading("Main areas", 2)
key_value_table([
    ("Dashboard", "Calendar, date countdown, recently opened scripts, and quick actions."),
    ("All scripts", "The script and folder library, including nested folder navigation and search."),
    ("Cells", "An Excel-like workbook with multiple sheets, rows, columns, typed cells, and copy/paste."),
    ("Backup center", "Full-vault ZIP export, restore, and browser storage usage information."),
    ("Password & security", "Password change, current-password validation, and optional lock-screen hint."),
], header=("Area", "Purpose"))

heading("2. Password, lock screen, hint, and theme", 1)
heading("Initial password", 2)
para("On first use, the default case-sensitive password is SCOAW2899. Avaton immediately derives and stores a salted password hash for verification. The active password is not stored as readable plain text.")
heading("Changing the password", 2)
number("Open Password & security from the bottom of the sidebar.")
number("Enter the current password.")
number("Enter and confirm a new case-sensitive password of at least four characters.")
number("Optionally enter a hint. The hint is visible from the locked screen.")
number("Select Update password.")
callout("Security warning", "The hint should never contain the actual password. Anyone who can open the lock screen can reveal the hint.", RED_LIGHT)
heading("What password data is stored", 2)
bullet("A random salt and a PBKDF2-derived password hash are stored in browser localStorage.")
bullet("The optional password hint is also stored in localStorage.")
bullet("The password, salt, hash, and hint are not included in Avaton ZIP backups.")
bullet("If browser site data is erased, the customized password settings are erased too. A fresh Avaton profile starts with the default password.")
heading("Lock behavior", 2)
para("Lock Avaton hides the workspace and returns to the password screen. The browser session remembers an unlocked state only for the current tab session. Closing the session or explicitly locking removes that session flag.")
heading("Theme", 2)
para("Light or dark mode is saved in localStorage for the current browser profile. Theme choice is not part of the backup.")

heading("3. Scripts, nested scripts, and folders", 1)
heading("Scripts", 2)
bullet("A script has a title, editable text body, color, timestamps, attachments, folder assignment, and optional nested scripts.")
bullet("Changes to the title and text are automatically saved locally after a short delay.")
bullet("Recently opened scripts appear on the Dashboard.")
bullet("Search matches script titles and body text.")
heading("Nested scripts", 2)
para("A script can contain other scripts. Deleting a parent script deletes that parent, all nested descendants, and every attachment belonging to that family. The deletion confirmation explains how many nested scripts are affected.")
heading("Folders", 2)
bullet("Folders can contain multiple top-level scripts.")
bullet("Folders can contain other folders.")
bullet("New scripts created while a folder is open are placed in that folder.")
bullet("An existing script can be moved with the Folder selector in the script editor.")
bullet("Folders and their hierarchy are included in a full backup.")
callout("Deletion is permanent", "Script deletion removes local script records and their stored file chunks. There is no recycle bin or undo. Create a backup before deleting anything important.", RED_LIGHT)

heading("4. Attachments: upload, storage, download, and preview", 1)
heading("Uploading files", 2)
para("Avaton accepts any file type through the browser file picker, including ZIP archives. Files are processed in 4 MB pieces and written to browser storage. The progress window shows the active filename, piece count, percentage, and transferred bytes.")
bullet("Cancel upload stops after the current 4 MB piece and removes incomplete data from that interrupted upload.")
bullet("There is no Avaton-defined upload size limit.")
bullet("The real limit is available disk space, browser storage quota, operating-system stability, and browser memory.")
heading("Downloading attachments", 2)
para("Avaton attempts to preserve the original filename and extension. If an old attachment has no extension, common types such as ZIP, PDF, JPG, PNG, MP4, and WebM receive a suitable extension during download.")
para("For chunked files in compatible Chrome or Edge versions, Avaton can write directly to a chosen destination. Older or less capable browsers may use a normal browser download.")
heading("Built-in viewers", 2)
key_value_table([
    ("Images", "JPG, JPEG, PNG, GIF, WebP, BMP, SVG, and AVIF where supported by the browser."),
    ("Video", "MP4, WebM, OGG, MOV, M4V, MKV, and AVI where the browser supports the file's codec."),
    ("PDF", "Displayed in the browser's embedded PDF viewer."),
    ("ZIP", "Displays the archive entry names; it does not extract or run archive contents."),
    ("Other types", "Stored and downloadable, but no built-in preview is provided."),
], header=("Type", "Viewer behavior"))
callout("Large video clarification", "Avaton does not impose a video preview size limit. However, the browser must still assemble access to the stored chunks and support the video's codec. Extremely large or high-bitrate videos can be slow, use significant resources, or fail on low-memory systems.", GOLD_LIGHT)
para("For non-video files, previews above 750 MB are blocked to reduce the chance of freezing the laptop. The file remains available for download.")

heading("5. Cells: Excel-like workbook", 1)
para("Cells is a local workbook inside Avaton. It is designed for planning, lists, dates, checklists, and pasted spreadsheet data. It is not Microsoft Excel and does not create an .xlsx file.")
heading("Workbook features", 2)
bullet("Multiple named sheets with sheet tabs.")
bullet("Add and delete sheets. At least one sheet must remain.")
bullet("Rename the active sheet.")
bullet("Add rows and columns.")
bullet("Delete the currently selected row or column. At least one row and column must remain.")
bullet("Editable column names and visible Excel-style column letters and row numbers.")
bullet("Text, number, date, and checkbox column types.")
bullet("Active-cell address and formula-style value bar.")
bullet("Arrow, Enter, and Tab navigation.")
bullet("Paste tabular rows copied from Excel or another spreadsheet.")
heading("Cells limitations", 2)
bullet("No formulas or formula calculation engine.")
bullet("No charts, pivot tables, merged cells, formatting rules, sorting, filters, or Excel file import/export.")
bullet("Column type applies to the whole column, not to individual cells.")
bullet("Cells is backed up as structured Avaton data, not as a separate Excel workbook.")

heading("6. Calendar and countdowns", 1)
para("The calendar defaults to year 2026 and supports earlier or later years with navigation controls. Every visible date shows its relationship to the current local date: days left, today, or days ago. Selecting a date opens a larger countdown summary.")
bullet("Countdowns use the computer's current local date and timezone.")
bullet("Changing the computer clock or timezone changes the displayed countdown.")
bullet("Calendar selections are not saved as permanent records or included in backup.")

heading("7. Exactly where Avaton stores data", 1)
callout("Working data is not stored beside index.html", "Opening the Avaton files does not create a visible data folder next to the app. The browser stores the working copy inside its private site-storage area.", BLUE_LIGHT)
key_value_table([
    ("IndexedDB database", "avaton-vault"),
    ("Scripts store", "Titles, text, colors, nesting, folder links, timestamps, and attachment metadata."),
    ("Folders store", "Folder names and parent-folder relationships."),
    ("Cells store", "The complete multi-sheet Cells workbook."),
    ("File chunks store", "Attachment bytes divided into 4 MB pieces."),
    ("Legacy files store", "Compatibility storage for attachments created by earlier builds."),
    ("localStorage", "Password hash, password salt, hint, and theme."),
    ("sessionStorage", "Temporary unlocked-session flag."),
], header=("Storage location", "What it contains"))
heading("What 'tied to this browser' means", 2)
bullet("Chrome and Edge profiles have separate storage.")
bullet("Normal and private/incognito windows have separate or temporary storage.")
bullet("A different address or port may be treated as a different site and therefore a different vault.")
bullet("Clearing site data, deleting the browser profile, resetting the browser, or reinstalling Windows can remove the working copy.")
bullet("Copying only index.html, styles.css, and app.js does not copy the saved vault.")
heading("Storage capacity", 2)
para("Backup Center displays the browser's estimated usage and quota. The quota is controlled by the browser and available device storage. It is not a guaranteed reservation. A file can still fail if the disk fills, the browser changes its quota, or the operating system interrupts the write.")

heading("8. How the backup system works", 1)
heading("Full backup", 2)
para("Export everything creates an Avaton ZIP backup. The ZIP contains an avaton-manifest.json file plus the original attachment bytes arranged under file paths. The manifest includes all scripts, nesting, folders, every Cells sheet, row and column definitions, values, checkboxes, attachment metadata, and timestamps.")
heading("Single-script backup", 2)
para("Exporting one script includes that script and all of its nested descendants. The current implementation also places the complete Cells workbook in that ZIP. Folder metadata is limited to folders directly referenced by the exported scripts.")
heading("Compression and file size", 2)
bullet("The backup is a standard ZIP container.")
bullet("When the browser supports the required streaming compression API, ZIP entries are compressed.")
bullet("If the browser does not support that compression API, the result is still a ZIP but entries may be stored without compression.")
bullet("ZIP, MP4, JPEG, and many PDFs are already compressed and may shrink very little.")
bullet("Text and some office-style data often compress substantially.")
bullet("The current ZIP implementation does not support ZIP64. A single backup must remain below 4 GB.")
heading("Small versus large backup writing", 2)
bullet("Smaller backups are assembled as a browser download.")
bullet("Large backups use the browser's save-file picker when supported, allowing direct writing to a selected disk or USB destination.")
bullet("If direct file writing is unavailable, very large backups may require substantial memory or may fail.")
heading("What backup does to current files", 2)
callout("Export is non-destructive", "Creating a backup does not move, delete, rename, or change the current scripts, folders, Cells sheets, or attachments. It only reads the current vault and creates a separate ZIP copy.", GREEN_LIGHT)
heading("What is not included", 2)
bullet("Password, password salt, password hash, and password hint.")
bullet("Light/dark theme preference.")
bullet("Current unlocked session.")
bullet("Calendar's currently selected date.")
bullet("Browser permissions and storage quota.")

heading("9. What happens when a backup is restored", 1)
heading("Restore is a merge, not a factory reset", 2)
para("Importing a backup does not automatically erase the current vault first. Avaton reads the backup and writes its contents into the current browser database.")
key_value_table([
    ("Matching script ID", "The backup version replaces/updates the current script record."),
    ("New script ID", "The script is added."),
    ("Current script not in backup", "It remains in the vault."),
    ("Matching folder ID", "The backup folder record replaces/updates the current record."),
    ("New folder ID", "The folder is added."),
    ("Current folder not in backup", "It remains."),
    ("Attachment with matching ID", "Stored chunks for that attachment are replaced with the backup bytes."),
    ("Current attachment not referenced by imported scripts", "It generally remains with its existing script."),
    ("Cells workbook", "The backup's entire Cells workbook replaces the current Cells workbook."),
    ("Password/theme", "They remain unchanged because they are not imported."),
], header=("Item", "Result after restore"))
heading("Restore safety limitations", 2)
bullet("Restore is not wrapped in one all-or-nothing database transaction.")
bullet("If the browser, laptop, USB drive, or tab stops midway, some items may already have been restored while later items are missing.")
bullet("Importing a corrupted or incomplete ZIP reports an error, but it may not undo earlier writes from that attempt.")
bullet("There is no automatic duplicate-resolution screen or preview-before-merge.")
callout("Safest restore practice", "Before importing into a vault that already contains important work, first export the current vault to a separate backup file. Then import the other backup.", GOLD_LIGHT)

heading("10. Limitations", 1)
key_value_table([
    ("Local-only architecture", "No cloud sync, server account, remote recovery, or automatic off-device copy."),
    ("Browser dependency", "Quota, codec support, PDF support, save picker, compression, and performance vary by browser."),
    ("Backup maximum", "No ZIP64; one ZIP backup must remain below 4 GB."),
    ("Very large files", "May be slow, may exhaust disk or memory, and can make preview/backup operations resource-intensive."),
    ("Video playback", "Any size is allowed by Avaton, but browser codec support and hardware resources still apply."),
    ("Non-video preview", "Files over 750 MB are not previewed."),
    ("ZIP viewer", "Lists filenames only; no extraction, editing, password-protected ZIP support, or execution."),
    ("Cells", "No formulas, charts, Excel file format, advanced formatting, sorting, or filtering."),
    ("Password recovery", "No reset link or recovery key. The password itself is not in the backup."),
    ("Encryption", "The lock screen verifies access in the app, but stored IndexedDB files and ZIP backups are not end-to-end encrypted by Avaton."),
    ("Deletion", "No recycle bin or undo for deleted scripts and attachments."),
    ("Restore", "Merge-based and not transactional; current unmatched data remains."),
], header=("Limitation", "Practical meaning"))
callout("Privacy versus encryption", "Avaton keeps data local and does not upload it to a server. Local-only is not the same as encrypted storage. Someone with access to the browser profile, computer account, or unencrypted backup ZIP may be able to obtain the underlying data.", RED_LIGHT)

heading("11. Recommended backup and recovery workflow", 1)
heading("Routine protection", 2)
number("Create a full ZIP backup after important changes or large uploads.")
number("Wait for the success message before closing Avaton or disconnecting a USB drive.")
number("Confirm the ZIP exists and its size is plausible.")
number("Keep two copies: one on a USB/external drive and one in another safe location.")
number("Periodically test restoration in a separate browser profile or separate Avaton site address.")
heading("Before reinstalling or deleting the laptop", 2)
number("Export everything.")
number("Copy the ZIP to a USB drive.")
number("Open the USB drive and verify the file is present.")
number("If practical, copy the ZIP to a second device and test that it can be selected for restore.")
number("Only then clear browser data, reinstall Windows, or erase the laptop.")
heading("After reinstalling", 2)
number("Place the Avaton HTML, CSS, and JavaScript files together.")
number("Open Avaton using a stable address in a current Chrome or Edge browser.")
number("The fresh profile initially uses the default password SCOAW2899.")
number("Open Backup Center and choose the saved ZIP backup.")
number("Wait for restoration to finish. Do not disconnect the USB or close the browser.")
number("Confirm scripts, nested scripts, folders, every Cells sheet, and several attachments.")
number("Change the password and hint again because password settings are not restored from backup.")
heading("If a restore fails", 2)
bullet("Keep the original backup unchanged.")
bullet("Check free disk space and browser storage quota.")
bullet("Try a current Chrome or Edge version.")
bullet("Copy the backup from USB to the laptop before restoring if the USB connection is unstable.")
bullet("Do not repeatedly import into the only valuable working copy without first backing it up.")

heading("12. Quick answers", 1)
key_value_table([
    ("Does backup delete my current files?", "No. Export only creates a separate copy."),
    ("Does restore delete my current files?", "No automatic wipe. It merges records; matching IDs update, unmatched current items remain. Cells is replaced as a whole."),
    ("Can I store a 1 GB file?", "Avaton allows it, but success depends on browser quota, free disk space, and system resources."),
    ("Can I upload ZIP files?", "Yes."),
    ("Can I preview any video size?", "Avaton does not impose a video-size limit, but codec and hardware limitations still apply."),
    ("Are all Cells sheets backed up?", "Yes, the full workbook is included."),
    ("Is my password backed up?", "No."),
    ("Will copying the app files copy my data?", "No. Export a ZIP backup."),
    ("Is the backup encrypted?", "No. Protect the USB or backup location appropriately."),
    ("Can one backup exceed 4 GB?", "Not with the current ZIP implementation. Export scripts separately to keep each ZIP below 4 GB."),
], header=("Question", "Answer"))

doc.add_paragraph()
callout("Final recommendation", "Treat the browser vault as the working copy and the ZIP as the recovery copy. A backup is only trustworthy after it has completed, exists on another device, and has been tested at least once.", GREEN_LIGHT)

doc.save(OUT)
print(OUT)
